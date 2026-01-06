"""
Pleadings Assistant - Layer 3

Generates professional legal pleadings and court documents.
Supports common document types with Ghana-specific formatting.

Features:
- Multiple document templates (Summons, Statement of Claim, Defence, etc.)
- LLM-powered content generation
- Ghana court formatting rules
- Document assembly and export
- Legal citation integration
"""

from typing import Dict, List, Optional, Any, Tuple
from dataclasses import dataclass, field, asdict
from datetime import datetime, timedelta
from enum import Enum
import json

from reasoning.llm_integration import (
    get_llm_orchestrator,
    TaskType,
    LLMResponse
)
from intelligence.citation_network import get_citation_network


class PleadingType(Enum):
    """Types of pleadings and court documents"""
    SUMMONS = "summons"
    STATEMENT_OF_CLAIM = "statement_of_claim"
    DEFENCE = "defence"
    COUNTERCLAIM = "counterclaim"
    REPLY = "reply"
    APPLICATION = "application"
    AFFIDAVIT = "affidavit"
    MEMORIAL = "memorial"
    NOTICE_OF_MOTION = "notice_of_motion"
    STATUTORY_DECLARATION = "statutory_declaration"


class CourtType(Enum):
    """Ghana court types"""
    SUPREME_COURT = "supreme_court"
    COURT_OF_APPEAL = "court_of_appeal"
    HIGH_COURT = "high_court"
    CIRCUIT_COURT = "circuit_court"
    DISTRICT_COURT = "district_court"
    CUSTOMARY_COURT = "customary_court"


@dataclass
class Party:
    """Party to litigation"""
    name: str
    capacity: str  # "Plaintiff/Appellant", "Defendant/Respondent", "Interested Party"
    address: str
    lawyer: Optional[str] = None
    lawyer_address: Optional[str] = None
    reference: Optional[str] = None


@dataclass
class ClaimDetail:
    """Individual claim or allegation"""
    description: str
    legal_basis: str  # Statute or common law principle
    amount: Optional[float] = None
    relief_type: str = "general"  # "damages", "specific performance", "injunction", "declaratory"


@dataclass
class PleadingMetadata:
    """Metadata for pleading document"""
    case_name: str
    case_number: str
    court: CourtType
    filing_date: str
    plaintiff: Party
    defendant: Party
    judge_assigned: Optional[str] = None
    previous_actions: Optional[List[str]] = field(default_factory=list)
    limitation_period_expires: Optional[str] = None


@dataclass
class Pleading:
    """Generated pleading document"""
    pleading_type: PleadingType
    metadata: PleadingMetadata
    
    # Document sections
    caption: str  # Court and case header
    title: str  # Document title
    preamble: str  # Introductory language
    facts: List[str]  # Numbered factual paragraphs
    legal_contentions: List[str]  # Legal arguments (numbered)
    relief_sought: List[str]  # Specific relief requested
    
    # Closing sections
    verification: Optional[str] = None  # Sworn statement
    conclusion: str = ""
    signatures: List[Dict[str, str]] = field(default_factory=list)
    
    # Metadata
    generated_date: str = field(default_factory=lambda: datetime.now().isoformat())
    version: str = "1.0"
    generation_tokens: int = 0
    generation_cost: float = 0.0
    citations: List[Dict[str, Any]] = field(default_factory=list)

    def to_text(self) -> str:
        """Convert to plain text format"""
        doc = f"""{self.caption}

{self.title}

{self.preamble}

FACTS AND BACKGROUND
"""
        for i, fact in enumerate(self.facts, 1):
            doc += f"{i}. {fact}\n\n"
        
        doc += "\nLEGAL CONTENTIONS\n"
        for i, contention in enumerate(self.legal_contentions, 1):
            doc += f"{i}. {contention}\n\n"
        
        doc += "\nRELIEF SOUGHT\n"
        for relief in self.relief_sought:
            doc += f"• {relief}\n"
        
        if self.verification:
            doc += f"\n\nVERIFICATION\n{self.verification}\n"
        
        if self.conclusion:
            doc += f"\n\n{self.conclusion}\n"
        
        doc += "\n\nDated this _____ day of _________ 20____\n\n"
        for sig in self.signatures:
            doc += f"____________________\n{sig.get('name', '')}\n"
            if sig.get('title'):
                doc += f"{sig.get('title')}\n"
            doc += "\n\n"
        
        return doc

    def to_docx(self) -> bytes:
        """Export to Word document format"""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Add caption and title
            caption_para = doc.add_paragraph(self.caption)
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_para.runs[0].bold = True
            
            title_para = doc.add_paragraph(self.title)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.runs[0].bold = True
            
            doc.add_paragraph()
            doc.add_paragraph(self.preamble)
            
            # Facts section
            doc.add_heading('FACTS AND BACKGROUND', level=1)
            for i, fact in enumerate(self.facts, 1):
                doc.add_paragraph(fact, style=f'List Number {i}')
            
            # Legal contentions
            doc.add_heading('LEGAL CONTENTIONS', level=1)
            for i, contention in enumerate(self.legal_contentions, 1):
                doc.add_paragraph(contention, style=f'List Number {i}')
            
            # Relief sought
            doc.add_heading('RELIEF SOUGHT', level=1)
            for relief in self.relief_sought:
                doc.add_paragraph(relief, style='List Bullet')
            
            # Citations if any
            if self.citations:
                doc.add_heading('AUTHORITIES CITED', level=1)
                for citation in self.citations:
                    doc.add_paragraph(f"{citation.get('name', '')}: {citation.get('citation', '')}")
            
            # Save to bytes
            import io
            output = io.BytesIO()
            doc.save(output)
            return output.getvalue()
        
        except ImportError:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary"""
        return asdict(self)

    def to_json(self) -> str:
        """Convert to JSON"""
        return json.dumps(self.to_dict(), indent=2, default=str)


class PleadingsAssistant:
    """Generate professional legal pleadings"""

    def __init__(self):
        self.llm = get_llm_orchestrator()
        self.citation_network = get_citation_network()
        self.court_formats = self._load_court_formats()
        self.templates = self._load_templates()

    def _load_court_formats(self) -> Dict[CourtType, Dict[str, str]]:
        """Load court-specific formatting rules"""
        return {
            CourtType.SUPREME_COURT: {
                "name": "IN THE SUPREME COURT OF GHANA",
                "caption_style": "IN THE SUPREME COURT OF GHANA\nHOLDING BOTH ORIGINAL AND APPELLATE JURISDICTION",
                "case_format": "[YEAR] GHASC [NUMBER]",
                "filing_rules": "Must include bench constitution",
                "page_style": "A4, 1.5 line spacing"
            },
            CourtType.COURT_OF_APPEAL: {
                "name": "IN THE COURT OF APPEAL OF GHANA",
                "caption_style": "IN THE COURT OF APPEAL OF GHANA",
                "case_format": "[YEAR] SCGLR/CA [NUMBER]",
                "filing_rules": "Appellant and Respondent parties only",
                "page_style": "A4, 1.5 line spacing"
            },
            CourtType.HIGH_COURT: {
                "name": "IN THE HIGH COURT OF JUSTICE",
                "caption_style": "IN THE HIGH COURT OF JUSTICE, ACCRA",
                "case_format": "HC/[YEAR]/[NUMBER]",
                "filing_rules": "Standard civil procedure applies",
                "page_style": "A4, 1.5 line spacing"
            },
            CourtType.CIRCUIT_COURT: {
                "name": "IN THE CIRCUIT COURT",
                "caption_style": "IN THE CIRCUIT COURT AT [LOCATION]",
                "case_format": "CC/[YEAR]/[NUMBER]",
                "filing_rules": "Limited jurisdiction rules apply",
                "page_style": "A4, single spacing"
            }
        }

    def _load_templates(self) -> Dict[PleadingType, str]:
        """Load pleading templates"""
        return {
            PleadingType.SUMMONS: """TO THE ABOVE NAMED DEFENDANT:

This Summons is issued against you by the above named Plaintiff(s).

You are required to ENTER an APPEARANCE within {appearance_period} days from the service of this Summons on you, failing which the Plaintiff(s) may proceed and judgment may be given against you without further notice.

If you intend to defend this action, you must file and serve a DEFENCE on the Plaintiff(s) or the Plaintiff's Solicitors within {defence_period} days after the expiry of the period within which you were required to enter an appearance.

TAKE NOTE: If you do not enter an appearance or file a Defence, judgment may be entered against you by default.

Dated at {location} this _____ day of _________ 20____

                                        COURT SEAL/STAMP

{solicitor_details}""",

            PleadingType.STATEMENT_OF_CLAIM: """STATEMENT OF CLAIM

The Plaintiff/Appellant claims against the Defendant/Respondent as follows:

[FACTS]

PARTICULARS OF CLAIM

[ALLEGATIONS]

[RELIEF SOUGHT]

And the Plaintiff claims accordingly.""",

            PleadingType.DEFENCE: """DEFENCE

The Defendant/Respondent hereby makes the following defence to the Statement of Claim:

1. The Defendant admits the allegations in paragraphs [NUMBERS] of the Statement of Claim.

2. The Defendant denies the allegations in paragraphs [NUMBERS] of the Statement of Claim.

3. As to the remaining allegations, the Defendant puts the Plaintiff to the proof thereof.

4. [SPECIFIC DEFENCES]

And the Defendant contends that the Plaintiff has not proved the case and claims that the action be dismissed.""",

            PleadingType.AFFIDAVIT: """AFFIDAVIT

I, {deponent_name}, of {deponent_address}, [occupation], make oath and state as follows:

{numbered_statements}

AND I make this solemn affidavit conscientiously believing the same to be true and by virtue of the Statutory Declarations Act, 1972 (N.R.C.D. 110).

                                    ________________________
                                    [DEPONENT SIGNATURE]

Dated at {location} this _____ day of _________ 20____

[SWORN BEFORE ME]
                                    ________________________
                                    [COMMISSIONER NAME]
                                    Commissioner for Oaths""",

            PleadingType.APPLICATION: """IN THE HIGH COURT OF JUSTICE
APPLICATION FOR [RELIEF SOUGHT]

Notice of Motion
Take notice that the applicant {applicant_name} will apply to this Honourable Court on [DATE] at [TIME] or such other date as Counsel may be heard for the following relief:

[APPLICATION RELIEF]

Grounds:
{grounds}

Dated at {location} this _____ day of _________ 20____

                                        {solicitor_name}
                                        Applicant's Counsel"""
        }

    def generate_pleading(
        self,
        pleading_type: PleadingType,
        metadata: PleadingMetadata,
        facts: List[str],
        legal_basis: List[str],
        relief_sought: List[str],
        use_llm: bool = True
    ) -> Pleading:
        """Generate a pleading document"""
        
        court_format = self.court_formats[metadata.court]
        
        # Build caption
        caption = f"""{court_format['caption_style']}

CASE NO. {metadata.case_number}

BETWEEN

{metadata.plaintiff.name.upper()}
(Plaintiff/Appellant)
-and-

{metadata.defendant.name.upper()}
(Defendant/Respondent)"""
        
        # Build title
        title_map = {
            PleadingType.SUMMONS: "SUMMONS",
            PleadingType.STATEMENT_OF_CLAIM: "STATEMENT OF CLAIM",
            PleadingType.DEFENCE: "DEFENCE",
            PleadingType.COUNTERCLAIM: "COUNTERCLAIM",
            PleadingType.REPLY: "REPLY",
            PleadingType.AFFIDAVIT: "AFFIDAVIT",
            PleadingType.APPLICATION: "APPLICATION FOR RELIEF"
        }
        title = title_map.get(pleading_type, "LEGAL PLEADING")
        
        # Build preamble
        preamble = f"""TO THE HONOURABLE COURT:

Please be advised that:

The Plaintiff/Applicant by their Solicitors, {metadata.plaintiff.lawyer or 'hereinafter identified'}, respectfully submits this {title} in the matter above."""
        
        # Generate detailed facts and contentions using LLM if needed
        facts_numbered = facts
        contentions_numbered = legal_basis
        
        if use_llm:
            # Use LLM to expand and refine facts
            facts_prompt = f"""Based on these facts, generate 5-7 detailed numbered factual paragraphs suitable for a {pleading_type.value.replace('_', ' ')} in Ghana courts:

Facts: {', '.join(facts)}

Format as numbered paragraphs (1. ... 2. ..., etc):"""
            
            try:
                facts_response = self.llm.generate(
                    facts_prompt,
                    task_type=TaskType.PLEADING_DRAFTING
                )
                # Parse LLM response into numbered facts
                facts_text = facts_response.text
                facts_numbered = [f.strip() for f in facts_text.split('\n') if f.strip() and any(c.isdigit() for c in f[:3])]
            except:
                facts_numbered = facts
        
        # Build pleading object
        pleading = Pleading(
            pleading_type=pleading_type,
            metadata=metadata,
            caption=caption,
            title=title,
            preamble=preamble,
            facts=facts_numbered,
            legal_contentions=contentions_numbered,
            relief_sought=relief_sought,
            conclusion=f"WHEREFORE the {self._get_party_title(pleading_type, True)} prays the Court for the above relief."
        )
        
        # Add citations if available
        all_authorities = []
        for contention in contentions_numbered:
            citations = self.citation_network.extract_citations(contention)
            for citation in citations:
                all_authorities.append({
                    'citation': citation.citation,
                    'name': citation.case_name,
                    'relationship': citation.relationship.value if citation.relationship else 'cited'
                })
        
        pleading.citations = all_authorities[:10]  # Top 10 citations
        
        return pleading

    def generate_summons(
        self,
        metadata: PleadingMetadata,
        relief: List[str]
    ) -> Pleading:
        """Generate a summons"""
        return self.generate_pleading(
            PleadingType.SUMMONS,
            metadata,
            facts=[],
            legal_basis=[],
            relief_sought=relief,
            use_llm=False
        )

    def generate_statement_of_claim(
        self,
        metadata: PleadingMetadata,
        facts: List[str],
        legal_basis: List[str],
        relief: List[str]
    ) -> Pleading:
        """Generate a statement of claim"""
        return self.generate_pleading(
            PleadingType.STATEMENT_OF_CLAIM,
            metadata,
            facts=facts,
            legal_basis=legal_basis,
            relief_sought=relief
        )

    def generate_defence(
        self,
        metadata: PleadingMetadata,
        admissions: List[str],
        denials: List[str],
        specific_defences: List[str]
    ) -> Pleading:
        """Generate a defence"""
        legal_basis = (
            [f"Admits: {a}" for a in admissions] +
            [f"Denies: {d}" for d in denials] +
            specific_defences
        )
        return self.generate_pleading(
            PleadingType.DEFENCE,
            metadata,
            facts=[],
            legal_basis=legal_basis,
            relief_sought=["Dismissal of the action"]
        )

    def generate_affidavit(
        self,
        metadata: PleadingMetadata,
        deponent_name: str,
        deponent_address: str,
        statements: List[str],
        sworn_before: str = "Commissioner for Oaths"
    ) -> Pleading:
        """Generate an affidavit"""
        pleading = self.generate_pleading(
            PleadingType.AFFIDAVIT,
            metadata,
            facts=statements,
            legal_basis=[],
            relief_sought=[]
        )
        
        # Customize for affidavit
        numbered_statements = "\n\n".join([f"{i}. {s}" for i, s in enumerate(statements, 1)])
        pleading.verification = f"""I, {deponent_name}, of {deponent_address}, make oath and state as follows:

{numbered_statements}

AND I make this solemn affidavit conscientiously believing the same to be true."""
        
        pleading.signatures = [
            {
                'name': deponent_name,
                'title': 'Deponent'
            },
            {
                'name': sworn_before,
                'title': 'Commissioner for Oaths'
            }
        ]
        
        return pleading

    def _get_party_title(self, pleading_type: PleadingType, is_plaintiff: bool = True) -> str:
        """Get appropriate party title"""
        if pleading_type in [PleadingType.STATEMENT_OF_CLAIM, PleadingType.SUMMONS]:
            return "Plaintiff" if is_plaintiff else "Defendant"
        elif pleading_type in [PleadingType.DEFENCE, PleadingType.COUNTERCLAIM]:
            return "Defendant" if is_plaintiff else "Plaintiff"
        else:
            return "Applicant" if is_plaintiff else "Respondent"

    def batch_generate(
        self,
        pleading_requests: List[Dict[str, Any]]
    ) -> List[Pleading]:
        """Generate multiple pleadings"""
        pleadings = []
        for request in pleading_requests:
            try:
                pleading_type = PleadingType[request.get('type', 'STATEMENT_OF_CLAIM').upper()]
                pleading = self.generate_pleading(
                    pleading_type=pleading_type,
                    metadata=request.get('metadata'),
                    facts=request.get('facts', []),
                    legal_basis=request.get('legal_basis', []),
                    relief_sought=request.get('relief', [])
                )
                pleadings.append(pleading)
            except Exception as e:
                print(f"Error generating pleading: {e}")
                continue
        
        return pleadings


# Global instance
_assistant: Optional[PleadingsAssistant] = None


def get_pleadings_assistant() -> PleadingsAssistant:
    """Get or create pleadings assistant singleton"""
    global _assistant
    if _assistant is None:
        _assistant = PleadingsAssistant()
    return _assistant


if __name__ == "__main__":
    # Test pleadings generation
    assistant = get_pleadings_assistant()
    print("Pleadings Assistant ready")
